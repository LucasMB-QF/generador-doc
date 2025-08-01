from fastapi import FastAPI, UploadFile, File, Request, HTTPException
from fastapi.responses import Response, HTMLResponse
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from openpyxl import load_workbook
from docx import Document
from docx.text.paragraph import Paragraph
import re
from io import BytesIO
from pathlib import Path
import logging
from urllib.parse import quote

# Configuración básica de logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Rutas
current_dir = Path(__file__).resolve().parent
templates_dir = current_dir / "templates"

app = FastAPI()

# Archivos estáticos
app.mount("/static", StaticFiles(directory="templates"), name="static")

# Templates Jinja2
templates = Jinja2Templates(directory=str(templates_dir))

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Regex para {{Hoja!Celda}} o {{Hoja!Rango}}
campo_regex = re.compile(r"\{\{\s*([^\{\}]+?)\s*\}\}")

# --- Funciones de lectura desde Excel ---

def obtener_valor(wb, hoja_nombre, celda):
    try:
        hoja = wb[hoja_nombre]
        valor = hoja[celda].value
        return str(valor) if valor is not None else ""
    except Exception as e:
        logger.error(f"Error en celda {hoja_nombre}!{celda}: {str(e)}")
        return ""

def obtener_valores_rango(wb, hoja_nombre, rango):
    try:
        hoja = wb[hoja_nombre]
        celdas = hoja[rango]
        fila = celdas[0]
        return [str(c.value) if c.value is not None else "" for c in fila]
    except Exception as e:
        logger.error(f"Error en rango {hoja_nombre}!{rango}: {str(e)}")
        return []

# --- Reemplazo de campos en texto ---

def reemplazar_campos(texto, wb):
    def reemplazo(match):
        campo = match.group(1)
        if '!' in campo:
            hoja, celda_o_rango = campo.split('!', 1)
            hoja = hoja.strip()
            celda_o_rango = celda_o_rango.strip()
            if ':' in celda_o_rango:
                valores = obtener_valores_rango(wb, hoja, celda_o_rango)
                return ', '.join(valores)
            else:
                return obtener_valor(wb, hoja, celda_o_rango)
        return ""
    return campo_regex.sub(reemplazo, texto)

# Reemplazo dentro de párrafos manteniendo estilo
def reemplazar_en_parrafo(parrafo: Paragraph, wb):
    for run in parrafo.runs:
        if campo_regex.search(run.text):
            run.text = reemplazar_campos(run.text, wb)

# --- Procesamiento del documento Word ---

def procesar_documento(doc, wb):
    for p in doc.paragraphs:
        reemplazar_en_parrafo(p, wb)

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    reemplazar_en_parrafo(p, wb)

# --- Rutas FastAPI ---

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/procesar")
async def procesar(
    archivo_excel: UploadFile = File(...),
    archivo_word: UploadFile = File(...)
):
    try:
        logger.info("Iniciando procesamiento de archivos...")

        # Validaciones
        if not archivo_excel.filename.endswith(('.xlsx', '.xlsm')):
            raise HTTPException(400, "El archivo Excel debe ser .xlsx o .xlsm")
        if not archivo_word.filename.endswith('.docx'):
            raise HTTPException(400, "El archivo Word debe ser .docx")

        # Leer archivos
        excel_content = await archivo_excel.read()
        word_content = await archivo_word.read()

        # Procesar documentos
        with BytesIO(excel_content) as excel_stream:
            wb = load_workbook(filename=excel_stream, data_only=True)

            with BytesIO(word_content) as word_stream:
                doc = Document(word_stream)
                procesar_documento(doc, wb)

                output_stream = BytesIO()
                doc.save(output_stream)
                output_stream.seek(0)

                logger.info("Procesamiento completado correctamente")

                nombre_base = archivo_word.filename.rsplit(".", 1)[0]
                nombre_generado = f"{nombre_base} (generado).docx"
                nombre_generado_seguro = quote(nombre_generado)

                return Response(
                    content=output_stream.getvalue(),
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={
                        "Content-Disposition": f'attachment; filename="{nombre_generado_seguro}"',
                        "Access-Control-Expose-Headers": "Content-Disposition"
                    }
                )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error en el procesamiento: {str(e)}", exc_info=True)
        raise HTTPException(500, f"Error interno del servidor: {str(e)}")

# --- Página de error personalizada ---

@app.exception_handler(HTTPException)
async def http_exception_handler(request, exc):
    return templates.TemplateResponse(
        "error.html",
        {
            "request": request,
            "status_code": exc.status_code,
            "detail": exc.detail
        },
        status_code=exc.status_code
    )
